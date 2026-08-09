"""
Extract plain text from an uploaded CV (PDF / DOCX / Markdown / TXT).
"""

import io
import re
from typing import Tuple


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _from_docx(data: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _from_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_cv_text(filename: str, data: bytes) -> Tuple[str, str]:
    """
    Returns (text, error). On success error is "".
    """
    name = (filename or "").lower()
    try:
        if name.endswith(".pdf"):
            text = _from_pdf(data)
        elif name.endswith((".docx", ".doc")):
            text = _from_docx(data)
        elif name.endswith((".md", ".markdown", ".txt")):
            text = _from_text(data)
        else:
            return "", f"Unsupported file type: {filename}"
    except Exception as e:
        return "", f"Could not read {filename}: {e}"

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    if len(text) < 100:
        return text, "Very little text extracted — if this is a scanned PDF, try a Word or Markdown version."
    return text, ""


def guess_name(cv_text: str) -> str:
    """Best-effort: the first short non-empty line is usually the name."""
    for line in cv_text.splitlines()[:8]:
        line = line.strip().lstrip("#").strip()
        if not line:
            continue
        words = line.split()
        if 1 < len(words) <= 5 and not any(c.isdigit() for c in line) and "@" not in line:
            return line.title() if line.isupper() else line
    return ""


def suggest_keywords(cv_text: str, limit: int = 5) -> list:
    """Suggest search keywords from role titles found in the CV."""
    text = cv_text.lower()
    roles = [
        "ai project manager", "project manager", "program manager", "product manager",
        "delivery manager", "engineering manager", "technical lead", "team lead",
        "machine learning engineer", "ml engineer", "ai engineer", "data scientist",
        "data engineer", "data analyst", "software engineer", "backend engineer",
        "frontend engineer", "full stack engineer", "devops engineer", "cloud engineer",
        "nlp engineer", "computer vision engineer", "research scientist",
        "business analyst", "scrum master", "qa engineer", "solutions architect",
        "lecturer", "professor", "consultant", "designer", "accountant", "marketing manager",
    ]
    found = [r.title() for r in roles if r in text]
    return found[:limit] if found else ["Project Manager"]
